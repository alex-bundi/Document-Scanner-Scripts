import pathlib
import codecs
import docx
import zipfile
import PyPDF2

#x = {"filename": "file1.txt", "page_no": 1, "content": "Lorem ipsum dolor sit amet."}

class ReadContents: 

    def read_txt(self, files_src, files_present)-> list:
        """Only reads documents with the file extension .txt"""

        self.files_src = files_src
        self.files_present = files_present
        self.txt_lines = []

        for file_name, file_ext in self.files_present:
            if file_ext == ".txt":
                self.full_path = pathlib.PurePath(self.files_src, file_name) # To avoid FileNotFoundError
                with open(self.full_path, "r") as scanned_file:
                    self.content = scanned_file.read().replace("\n", "").lower()
                    self.txt_lines.append({"filename": file_name, "content": self.content})        
        return self.txt_lines
    
    def handle_unicodeescape(self, text):
        self.text = text
        return text.encode("utf-8").decode("unicode_escape")
    
    def read_docx(self, files_src, files_present)-> list:
        """Only reads .docx documents."""

        self.files_src = files_src
        self.files_present = files_present
        self.docx_lines = []
        
        for file_name, file_ext in self.files_present:
            
            if file_ext == ".docx":
                self.full_path = pathlib.PurePath(self.files_src, file_name) # To avoid FileNotFoundError

                try:
                    doc = docx.Document(self.full_path) # Read .docx files
                except zipfile.BadZipFile as e: 
                    """ 
                    To avoid "zipfile.BadZipFile: File is not a zip file" when a .docx file type is open,
                    as MS Word creates a temporary file that starts with ($)aff portal.
                    """
                    continue
                
                self.docx_info = {}
                self.docx_info["filename"] = file_name
                self.docx_info["content"] = []

                if len(doc.tables) == 0:
                    for paragraphs in doc.paragraphs:
                        self.docx_info["content"].append(self.handle_unicodeescape(paragraphs.text))

                if len(doc.tables) > 0:
                    for table in doc.tables: # Access cells within the table
                        for row in table.rows:
                            for cell in row.cells:
                                self.text = cell.text.replace('\n', '').strip()
                                self.docx_info["content"].append(self.text)

                self.docx_lines.append(self.docx_info)

        return self.docx_lines
    
    def read_pdf(self, files_src, files_present)-> list:
        """Only reads documents with the file extension .pdf"""

        self.files_src = files_src
        self.files_present = files_present
        self.pdf_lines = []

        for file_name, file_ext in self.files_present:       
            if file_ext == ".pdf":
                self.full_path = str(pathlib.PurePath(self.files_src, file_name)) # The conversion helps to access the pdf files easier

                with open(self.full_path, "rb") as pdf_file_obj:
                    self.pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
                    for page_number in range(self.pdf_reader.getNumPages()):
                        self.pdf_info = {}
                        self.pdf_info["filename"] = file_name
                        self.page_obj = self.pdf_reader.getPage(page_number)
                        self.page_text = self.page_obj.extractText().replace("\n", "")
                        self.page_no = f"Page {page_number + 1}"
                        self.pdf_info["page_no"] = self.page_no
                        self.pdf_info["content"] = self.handle_unicodeescape(self.page_text)

                        self.pdf_lines.append(self.pdf_info)

        return self.pdf_lines


